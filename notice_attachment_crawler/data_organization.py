import os
import re
import pandas as pd
import pdfplumber
import docx
from pathlib import Path
from fuzzywuzzy import process, fuzz
from tqdm import tqdm
import warnings
warnings.filterwarnings('ignore')

# ================= 1. 配置区 =================

INPUT_DIR = "D:\homework\football\notice_attachment_crawler\data"          # 你的2000+文件目录
OUTPUT_FILE = "D:\homework\football\notice_attachment_crawler"  # 最终输出
LOG_FILE = "D:\homework\football\notice_attachment_crawler"

# 通知类关键词（用于过滤垃圾文件）
NOTICE_KEYWORDS = [
    '关于授予', '关于公示', '关于发布', '关于申请', '通知', '公告',
    '特此通知', '各有关单位', '各申报单位', '经我协会审核', '现授予'
]

# 列名标准化映射：标准名 → [可能的原始列名变体]
# 这是核心，你需要根据2000+文件的实际列名持续扩展这个字典
COLUMN_MAP = {
    'name':       ['姓名', '名字', '运动员姓名', '选手姓名', 'Name'],
    'gender':     ['性别', '男/女', 'Gender'],
    'id_card':    ['身份证号', '身份证', '证件号码', '身份证号(身份证)', '公民身份号码', '证件号'],
    'level':      ['申请等级', '等级', '技术等级', '运动员等级', '授予等级', '申请等级', '等级标准'],
    'event':      ['赛事', '比赛', '竞赛名称', '赛事名称', '申请赛事', '运动会', '锦标赛'],
    'rank':       ['名次', '申请成绩', '成绩', '排名', '名次/成绩', '比赛名次', '申请成绩名次'],
    'team':       ['参赛队名称', '参赛队', '单位', '队伍', '参赛单位', '代表队', '参赛队名称', '所属单位'],
    'sport':      ['申请项目', '项目', '运动项目', '体育项目'],
    'ethnicity':  ['民族', '国籍'],
    'phone':      ['联系方式', '联系电话', '电话', '手机'],
    'birthplace': ['籍贯', '户籍', '出生地', '户口所在地'],
    'apply_date': ['申请时间', '申报时间', '报名时间'],
    'grant_date': ['授予时间', '备案时间', '授予日期', '公示时间'],
    'cert_no':    ['证书编号', '编号', '证书号'],
    'province':   ['省份', '省', '地区', '省市'],
    'year':       ['年份', '年度', 'Year'],
    'source_file':['来源文件', '文件名']
}

# 反向索引：原始列名 → 标准列名
VARIANT_TO_STD = {}
for std, variants in COLUMN_MAP.items():
    for v in variants:
        VARIANT_TO_STD[v.lower()] = std

# ================= 2. 文件分类器 =================

def is_notice_file(file_path: str, text_sample: str) -> bool:
    """
    判断是否为通知类文件。综合文件名+内容特征。
    如果文件包含明显的名单表格特征，即使文件名像通知也保留。
    """
    fname = Path(file_path).name
    
    # 文件名通知分
    name_score = sum(1 for k in NOTICE_KEYWORDS if k in fname)
    
    # 内容通知分（取前1500字符）
    text = str(text_sample)[:1500]
    content_score = sum(1 for k in NOTICE_KEYWORDS if k in text)
    
    # 名单特征分（如果有这些，大概率是名单）
    list_markers = ['序号', '姓名', '一级运动员', '二级运动员', '身份证号', '参赛队']
    list_score = sum(1 for m in list_markers if m in text)
    
    # 如果检测到表格/名单特征很强，不认为是通知
    if list_score >= 2:
        return False
    
    return (name_score * 2 + content_score) > 3

# ================= 3. 多格式提取器 =================

class ExcelExtractor:
    def extract(self, path):
        """读取所有sheet，返回DataFrame列表"""
        dfs = []
        try:
            xl = pd.ExcelFile(path)
            for sheet in xl.sheet_names:
                df = pd.read_excel(path, sheet_name=sheet, dtype=str)
                df = df.dropna(how='all').dropna(axis=1, how='all')
                # 过滤掉只有几行或没有姓名列的sheet
                if len(df) > 0 and len(df.columns) >= 3:
                    dfs.append(df)
        except Exception as e:
            print(f"[Excel Error] {path}: {e}")
        return dfs, ""

class PDFExtractor:
    def extract(self, path):
        """优先提取表格，同时返回全文用于分类"""
        dfs = []
        full_text = ""
        try:
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text() or ""
                    full_text += text + "\n"
                    
                    # 提取表格
                    tables = page.extract_tables()
                    for table in tables:
                        if table and len(table) > 1:
                            # 处理表头合并单元格的情况
                            header = [str(h).replace('\n', '') if h else f"col_{i}" 
                                     for i, h in enumerate(table[0])]
                            df = pd.DataFrame(table[1:], columns=header)
                            df = df.dropna(how='all').dropna(axis=1, how='all')
                            if len(df) > 0 and len(df.columns) >= 3:
                                dfs.append(df)
        except Exception as e:
            print(f"[PDF Error] {path}: {e}")
        return dfs, full_text

class WordExtractor:
    def extract(self, path):
        """提取Word中的表格和段落文本"""
        dfs = []
        full_text = ""
        try:
            doc = docx.Document(path)
            full_text = "\n".join([p.text for p in doc.paragraphs])
            
            # 提取Word内嵌表格
            for table in doc.tables:
                data = []
                for row in table.rows:
                    data.append([cell.text.strip() for cell in row.cells])
                if len(data) > 1:
                    df = pd.DataFrame(data[1:], columns=data[0])
                    df = df.dropna(how='all').dropna(axis=1, how='all')
                    if len(df.columns) >= 3:
                        dfs.append(df)
            
            # 对于纯文本名单（如附件2：郑宇桁 张晓宇...），尝试正则解析
            if not dfs:
                df = parse_word_text_list(full_text, path)
                if df is not None:
                    dfs.append(df)
        except Exception as e:
            print(f"[Word Error] {path}: {e}")
        return dfs, full_text

def parse_word_text_list(text, filename):
    """
    解析Word中的纯文本名单（如18.docx的格式）
    示例："郑宇桁 张晓宇 武浩伟..." 或带序号的列表
    """
    # 模式1：检测到"授予二级运动员称号的通知"后的姓名列表
    if '授予' in text and '运动员' in text:
        # 提取中文姓名（2-4个字），连续出现
        lines = text.split('\n')
        names = []
        for line in lines:
            # 去除空格，按连续中文字符提取
            line = line.strip()
            if not line or len(line) > 20:  # 跳过太长行（可能是正文）
                continue
            # 简单规则：如果是2-4个纯中文字符，认为是姓名
            if re.match(r'^[\u4e00-\u9fa5]{2,4}$', line.replace(' ', '')):
                names.append(line.replace(' ', ''))
        
        if names:
            df = pd.DataFrame({'name': names})
            df['level'] = '二级运动员' if '二级' in text else '一级运动员'
            df['sport'] = '足球'
            df['source_file'] = Path(filename).name
            # 从文件名或内容推断省份年份
            df['province'] = extract_province(filename)
            df['year'] = extract_year(filename, text)
            return df
    return None

# ================= 4. 列名标准化引擎 =================

def normalize_columns(df, filename):
    """
    将原始列名映射为标准列名。
    策略：精确匹配 → 模糊匹配(>85分) → 保留为ext_前缀
    """
    new_cols = {}
    reserved = []  # 记录本表独有列
    
    for col in df.columns:
        col_clean = str(col).strip().replace('\n', '').replace(' ', '').replace('\u3000', '')
        
        # 1. 精确匹配（忽略大小写）
        std = VARIANT_TO_STD.get(col_clean.lower())
        
        # 2. 模糊匹配
        if not std:
            all_variants = list(VARIANT_TO_STD.keys())
            if all_variants:
                match = process.extractOne(col_clean, all_variants, scorer=fuzz.ratio)
                if match and match[1] > 85:
                    std = VARIANT_TO_STD.get(match[0])
        
        if std:
            # 防止多列映射到同一标准列（取第一个）
            if std not in new_cols.values():
                new_cols[col] = std
            else:
                # 重复的标准列，保留为扩展列
                ext_name = f"ext_{col_clean}"
                new_cols[col] = ext_name
                reserved.append(ext_name)
        else:
            # 未匹配到标准列，保留原始名（加安全前缀）
            safe = f"ext_{col_clean}" if not col_clean.startswith('ext_') else col_clean
            new_cols[col] = safe
            reserved.append(safe)
    
    df = df.rename(columns=new_cols)
    
    # 注入元数据列
    df['source_file'] = Path(filename).name
    df['province'] = extract_province(filename)
    df['year'] = extract_year(filename, str(df.columns))
    
    return df, reserved

def extract_province(filename):
    """从文件名推断省份"""
    provinces = ['重庆', '江西', '江苏', '浙江', '厦门', '河北', '辽宁', '山西', 
                 '沈阳', '北京', '上海', '广东', '山东', '四川', '湖北']
    for p in provinces:
        if p in filename:
            return p
    return '未知'

def extract_year(filename, text_content=""):
    """从文件名或内容提取年份"""
    m = re.search(r'(20\d{2})', filename)
    if m:
        return int(m.group(1))
    # 尝试从内容找
    m2 = re.search(r'(20\d{2})', str(text_content)[:500])
    if m2:
        return int(m2.group(1))
    return None

# ================= 5. 主处理流水线 =================

def main():
    all_dfs = []
    logs = []  # 记录处理状态
    
    files = [f for f in Path(INPUT_DIR).rglob('*') 
             if f.suffix.lower() in ['.xls', '.xlsx', '.pdf', '.docx']]
    
    print(f"共发现 {len(files)} 个待处理文件...")
    
    for file_path in tqdm(files):
        fname = file_path.name
        status = "success"
        reason = ""
        extracted_dfs = []
        
        try:
            # 根据后缀选择提取器
            ext = file_path.suffix.lower()
            if ext in ['.xls', '.xlsx']:
                extractor = ExcelExtractor()
                dfs, text = extractor.extract(file_path)
            elif ext == '.pdf':
                extractor = PDFExtractor()
                dfs, text = extractor.extract(file_path)
            elif ext == '.docx':
                extractor = WordExtractor()
                dfs, text = extractor.extract(file_path)
            else:
                continue
            
            # 文件分类：通知类且无有效表格的，直接跳过
            if is_notice_file(str(file_path), text) and not dfs:
                status = "skipped"
                reason = "notice_file_no_list"
                logs.append({"file": fname, "status": status, "reason": reason, "rows": 0})
                continue
            
            # 过滤通知类PDF里的非名单表格（如盖章页、正文页）
            valid_dfs = []
            for df in dfs:
                cols_text = " ".join([str(c) for c in df.columns])
                # 必须包含"姓名"或"身份证号"或"等级"才认为是名单
                if any(k in cols_text for k in ['姓名', '名字', '身份证号', '等级', '序号']):
                    valid_dfs.append(df)
            
            if not valid_dfs:
                status = "skipped"
                reason = "no_valid_table"
                logs.append({"file": fname, "status": status, "reason": reason, "rows": 0})
                continue
            
            # 标准化每一张表
            for df in valid_dfs:
                if len(df) < 1:
                    continue
                df_norm, reserved = normalize_columns(df, fname)
                all_dfs.append(df_norm)
                
        except Exception as e:
            status = "error"
            reason = str(e)
            print(f"[Fatal Error] {fname}: {e}")
        
        logs.append({
            "file": fname, 
            "status": status, 
            "reason": reason, 
            "rows": len(valid_dfs) if 'valid_dfs' in dir() else 0
        })
    
    # ================= 6. 宽表合并 =================
    if not all_dfs:
        print("未提取到任何有效数据，请检查输入目录或列名映射！")
        return
    
    print(f"成功提取 {len(all_dfs)} 张表，开始合并...")
    
    # 使用pd.concat做宽表合并：相同列自动对齐，不同列保留为NaN
    combined = pd.concat(all_dfs, axis=0, ignore_index=True, sort=False)
    
    # 列顺序优化：核心列放前面，扩展列放后面
    core_cols = [c for c in ['name', 'gender', 'id_card', 'level', 'sport', 
                             'event', 'rank', 'team', 'province', 'year', 
                             'ethnicity', 'phone', 'birthplace', 'apply_date', 
                             'grant_date', 'cert_no', 'source_file'] 
                 if c in combined.columns]
    ext_cols = [c for c in combined.columns if c not in core_cols]
    combined = combined[core_cols + ext_cols]
    
    # 数据清洗
    if 'level' in combined.columns:
        combined['level'] = combined['level'].astype(str).str.replace('运动员', '').str.strip()
        combined['level'] = combined['level'].replace({
            '一级': '一级运动员', '二级': '二级运动员', '三级': '三级运动员'
        })
    
    if 'gender' in combined.columns:
        combined['gender'] = combined['gender'].astype(str).str.replace('男', '男').str.replace('女', '女')
        combined['gender'] = combined['gender'].apply(
            lambda x: x if x in ['男', '女'] else None
        )
    
    # 去重（基于关键字段）
    subset_cols = [c for c in ['name', 'id_card', 'level', 'event'] if c in combined.columns]
    if subset_cols:
        before = len(combined)
        combined = combined.drop_duplicates(subset=subset_cols, keep='first')
        print(f"去重：{before} → {len(combined)} 行")
    
    # 保存
    combined.to_excel(OUTPUT_FILE, index=False, engine='openpyxl')
    pd.DataFrame(logs).to_csv(LOG_FILE, index=False, encoding='utf-8-sig')
    
    print(f"\n 完成！")
    print(f"   合并后总行数: {len(combined)}")
    print(f"   总列数: {len(combined.columns)} (含核心列+各文件独有列)")
    print(f"   输出文件: {OUTPUT_FILE}")
    print(f"   处理日志: {LOG_FILE}")
    print(f"\n列名清单:")
    for i, col in enumerate(combined.columns, 1):
        print(f"   {i}. {col}")

if __name__ == "__main__":
    main()