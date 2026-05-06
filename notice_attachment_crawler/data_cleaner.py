#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
足球运动员技术等级名单清洗合并脚本
功能：
  1. 自动过滤通知类文件，只保留含运动员名单的文件
  2. 支持 Excel / Word / PDF 多格式提取
  3. 智能表头检测（处理标题行偏移、合并单元格等脏格式）
  4. 异构表格宽表合并：同义列统一，独有列保留为 ext_ 前缀
  5. 输出合并宽表 + 处理日志，方便审计
"""

import os
import re
import json
import pandas as pd
from pathlib import Path
from tqdm import tqdm
import warnings

warnings.filterwarnings('ignore')

# ========================== 1. 配置区 ==========================

INPUT_DIR = r"D:\homework\football\notice_attachment_crawler\data"   # 你的2000+文件目录
OUTPUT_DIR = r"D:\homework\football\notice_attachment_crawler\output"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "运动员等级合并宽表.xlsx")
LOG_FILE = os.path.join(OUTPUT_DIR, "处理日志.csv")
CHECKPOINT = os.path.join(OUTPUT_DIR, "checkpoint.json")

os.makedirs(OUTPUT_DIR, exist_ok=True)

# 通知类关键词（用于过滤垃圾文件）
NOTICE_KEYWORDS = [
    '关于授予', '关于公示', '关于发布', '关于申请', '特此通知',
    '各有关单位', '各申报单位', '各相关单位', '各会员协会',
    '经我协会审核', '现授予', '现公示', '公示期', '无异议'
]

# 列名标准化映射：标准名 → [可能的原始列名变体]
# 遇到2000+文件里有新列名，直接往这里加即可
COLUMN_MAP = {
    'name':       ['姓名', '名字', '运动员姓名', '选手姓名', 'Name'],
    'gender':     ['性别', '男/女', 'Gender'],
    'id_card':    ['身份证号', '身份证', '证件号码', '身份证号(身份证)',
                   '公民身份号码', '证件号', '身份证|'],
    'level':      ['申请等级', '等级', '技术等级', '运动员等级', '授予等级',
                   '等级标准', '申请等级|'],
    'event':      ['赛事', '比赛', '竞赛名称', '赛事名称', '申请赛事',
                   '运动会', '锦标赛', '联赛', '杯赛'],
    'rank':       ['名次', '申请成绩', '成绩', '排名', '名次/成绩',
                   '比赛名次', '申请成绩名次'],
    'team':       ['参赛队名称', '参赛队', '单位', '队伍', '参赛单位',
                   '代表队', '所属单位', '参赛队名称'],
    'sport':      ['申请项目', '项目', '运动项目', '体育项目'],
    'ethnicity':  ['民族', '国籍'],
    'phone':      ['联系方式', '联系电话', '电话', '手机'],
    'birthplace': ['籍贯', '户籍', '出生地', '户口所在地'],
    'apply_date': ['申请时间', '申报时间', '报名时间'],
    'grant_date': ['授予时间', '备案时间', '授予日期', '公示时间'],
    'cert_no':    ['证书编号', '编号', '证书号'],
    'province':   ['省份', '省', '地区', '省市'],
    'year':       ['年份', '年度', 'Year'],
    'source_file':['来源文件', '文件名']
}

# 构建反向索引：原始列名 → 标准列名
VARIANT_TO_STD = {}
for std, variants in COLUMN_MAP.items():
    for v in variants:
        VARIANT_TO_STD[v.lower().replace('|', '')] = std


# ========================== 2. 轻量模糊匹配器 ==========================

def simple_match(col_clean: str):
    """列名清洗后的匹配逻辑：精确 → 包含匹配"""
    col_lower = col_clean.lower().strip()
    # 精确匹配
    if col_lower in VARIANT_TO_STD:
        return VARIANT_TO_STD[col_lower]
    # 包含匹配（避免单字误匹配，要求长度>2）
    for v, std in VARIANT_TO_STD.items():
        if v in col_lower or col_lower in v:
            if len(v) > 2:
                return std
    return None


# ========================== 3. 文件分类器 ==========================

def is_notice_file(fname: str, text_sample: str) -> bool:
    """
    判断是否为通知类文件。
    策略：文件名+内容关键词计分，但如果检测到名单特征则保留。
    """
    name_score = sum(1 for k in NOTICE_KEYWORDS if k in fname)

    text = str(text_sample)[:1500]
    content_score = sum(1 for k in NOTICE_KEYWORDS if k in text)

    # 名单特征反证：只要有这些关键词，即使文件名像通知也保留
    list_markers = ['序号', '姓名', '一级运动员', '二级运动员',
                    '身份证号', '参赛队', '申请等级']
    list_score = sum(1 for m in list_markers if m in text)

    if list_score >= 2:
        return False

    return (name_score * 2 + content_score) > 3


# ========================== 4. 多格式提取器 ==========================

class ExcelExtractor:
    """智能Excel提取器：自动检测真实表头行，跳过标题行偏移"""

    def extract(self, path: Path):
        dfs = []
        try:
            xl = pd.ExcelFile(path)
            for sheet in xl.sheet_names:
                # 先不指定header，读取原始数据
                df_raw = pd.read_excel(path, sheet_name=sheet,
                                       dtype=str, header=None)
                df_raw = df_raw.dropna(how='all').dropna(axis=1, how='all')

                if len(df_raw) < 2:
                    continue

                # 检测真实表头行（前5行内找包含最多关键词的行）
                header_row = None
                for i in range(min(5, len(df_raw))):
                    row_text = " ".join([
                        str(x) for x in df_raw.iloc[i] if pd.notna(x)
                    ])
                    header_markers = [
                        '姓名', '序号', '等级', '性别', '身份证号',
                        '赛事', '项目', '申请'
                    ]
                    score = sum(1 for m in header_markers if m in row_text)
                    if score >= 3:
                        header_row = i
                        break

                if header_row is None:
                    header_row = 0  # fallback

                # 用检测到的行作为header重新读取
                df = pd.read_excel(path, sheet_name=sheet,
                                   dtype=str, header=header_row)
                df = df.dropna(how='all').dropna(axis=1, how='all')

                # 删除全是 Unnamed 的列
                unnamed = [c for c in df.columns
                           if str(c).startswith('Unnamed')]
                df = df.drop(columns=unnamed, errors='ignore')

                if len(df) > 0 and len(df.columns) >= 2:
                    dfs.append(df)

        except Exception as e:
            print(f"[Excel Error] {path.name}: {e}")
        return dfs, ""


class PDFExtractor:
    """PDF提取器：优先文字提取，扫描件需配合OCR"""

    def extract(self, path: Path):
        dfs = []
        full_text = ""

        # 文字型PDF提取
        try:
            from pdfminer.high_level import extract_text
            full_text = extract_text(str(path))
        except Exception:
            pass

        # 如果 pdfplumber 可用，尝试提取表格
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    tables = page.extract_tables()
                    for table in tables:
                        if table and len(table) > 1:
                            header = [
                                str(h).replace('\n', '') if h else f"col_{i}"
                                for i, h in enumerate(table[0])
                            ]
                            df = pd.DataFrame(table[1:], columns=header)
                            df = df.dropna(how='all').dropna(axis=1, how='all')
                            if len(df) > 0 and len(df.columns) >= 2:
                                dfs.append(df)
        except Exception:
            pass

        return dfs, full_text


class WordExtractor:
    """Word提取器：内嵌表格 + 纯文本名单解析"""

    def extract(self, path: Path):
        dfs = []
        full_text = ""
        try:
            import docx
            doc = docx.Document(path)
            full_text = "\n".join([p.text for p in doc.paragraphs])

            # 提取内嵌表格
            for table in doc.tables:
                data = []
                for row in table.rows:
                    data.append([cell.text.strip() for cell in row.cells])
                if len(data) > 1:
                    df = pd.DataFrame(data[1:], columns=data[0])
                    df = df.dropna(how='all').dropna(axis=1, how='all')
                    if len(df.columns) >= 2:
                        dfs.append(df)

            # 纯文本名单兜底
            if not dfs:
                df = parse_word_text_list(full_text, path.name)
                if df is not None:
                    dfs.append(df)

        except Exception as e:
            print(f"[Word Error] {path.name}: {e}")
        return dfs, full_text


def parse_word_text_list(text: str, filename: str):
    """
    解析Word中的纯文本名单（如附件2：郑宇桁 张晓宇...）
    """
    if not ('授予' in text and '运动员' in text):
        return None

    lines = text.split('\n')
    names = []
    for line in lines:
        line = line.strip()
        if not line or len(line) > 20:
            continue
        if re.match(r'^[\u4e00-\u9fa5]{2,4}$', line.replace(' ', '')):
            names.append(line.replace(' ', ''))

    if len(names) < 3:
        return None

    df = pd.DataFrame({'name': names})
    df['level'] = '二级运动员' if '二级' in text else '一级运动员'
    df['sport'] = '足球'
    df['source_file'] = filename
    df['province'] = extract_province(filename)
    df['year'] = extract_year(filename, text)
    return df


# ========================== 5. 辅助函数 ==========================

def extract_province(filename: str) -> str:
    provinces = [
        '重庆', '江西', '江苏', '浙江', '厦门', '河北', '辽宁', '山西',
        '沈阳', '北京', '上海', '广东', '山东', '四川', '湖北', '湖南',
        '安徽', '福建', '河南', '黑龙江', '吉林', '天津', '陕西'
    ]
    for p in provinces:
        if p in filename:
            return p
    return None


def extract_year(filename: str, text_content: str = "") -> int:
    m = re.search(r'(20\d{2})', filename)
    if m:
        return int(m.group(1))
    m2 = re.search(r'(20\d{2})', str(text_content)[:500])
    if m2:
        return int(m2.group(1))
    return None


def normalize_columns(df: pd.DataFrame, filename: str):
    """
    列名标准化：同义列合并，独有列保留为 ext_ 前缀
    """
    new_cols = {}
    for col in df.columns:
        col_clean = str(col).strip() \
            .replace('\n', '').replace(' ', '').replace('\u3000', '')
        std = simple_match(col_clean)

        if std:
            # 防止多列映射到同一标准列
            if std not in new_cols.values():
                new_cols[col] = std
            else:
                new_cols[col] = f"ext_{col_clean}"
        else:
            safe = f"ext_{col_clean}" \
                if not col_clean.startswith('ext_') else col_clean
            new_cols[col] = safe

    df = df.rename(columns=new_cols)

    # 注入元数据列
    df['source_file'] = filename
    df['province'] = extract_province(filename)
    df['year'] = extract_year(filename, str(df.columns))

    return df


# ========================== 6. 断点续传 ==========================

def load_checkpoint():
    if os.path.exists(CHECKPOINT):
        with open(CHECKPOINT, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}


def save_checkpoint(data):
    with open(CHECKPOINT, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


# ========================== 7. 主处理流水线 ==========================

def main():
    all_dfs = []
    logs = []
    checkpoint = load_checkpoint()

    # 扫描所有文件
    files = sorted([
        f for f in Path(INPUT_DIR).rglob('*')
        if f.suffix.lower() in ['.xls', '.xlsx', '.pdf', '.docx']
    ])

    print(f"共发现 {len(files)} 个待处理文件")

    for idx, file_path in enumerate(files, 1):
        fname = file_path.name
        fpath_str = str(file_path)

        # 断点：已处理且文件未修改则跳过
        if fname in checkpoint:
            print(f"[{idx}/{len(files)}] {fname} -> 已处理，跳过")
            continue

        print(f"[{idx}/{len(files)}] {fname}")
        status = "success"
        reason = ""
        valid_dfs = []

        try:
            ext = file_path.suffix.lower()

            if ext in ['.xls', '.xlsx']:
                extractor = ExcelExtractor()
                dfs, text = extractor.extract(file_path)

            elif ext == '.pdf':
                extractor = PDFExtractor()
                dfs, text = extractor.extract(file_path)
                # 文字型PDF兜底：尝试解析纯文本名单
                if not dfs:
                    df_text = parse_word_text_list(text, fname)
                    if df_text is not None:
                        dfs.append(df_text)

            elif ext == '.docx':
                extractor = WordExtractor()
                dfs, text = extractor.extract(file_path)

            else:
                continue

            # 通知过滤
            if is_notice_file(fname, text) and not dfs:
                status = "skipped"
                reason = "notice_file_no_list"
                logs.append({
                    "file": fname, "status": status,
                    "reason": reason, "rows": 0
                })
                checkpoint[fname] = {"status": status, "reason": reason}
                continue

            # 名单有效性校验
            for df in dfs:
                cols_text = " ".join([str(c) for c in df.columns])
                if any(k in cols_text for k in [
                    '姓名', '等级', '序号', '项目',
                    '名次', '参赛', 'name', 'level'
                ]):
                    valid_dfs.append(df)

            if not valid_dfs:
                status = "skipped"
                reason = "no_valid_table"
                logs.append({
                    "file": fname, "status": status,
                    "reason": reason, "rows": 0
                })
                checkpoint[fname] = {"status": status, "reason": reason}
                continue

            # 标准化 & 收集
            for df in valid_dfs:
                if len(df) < 1:
                    continue
                df_norm = normalize_columns(df, fname)
                all_dfs.append(df_norm)

        except Exception as e:
            status = "error"
            reason = str(e)
            print(f"  [Fatal Error] {e}")

        logs.append({
            "file": fname,
            "status": status,
            "reason": reason,
            "rows": sum(len(d) for d in valid_dfs) if valid_dfs else 0
        })
        checkpoint[fname] = {"status": status, "reason": reason}

        # 每10个文件保存一次断点
        if idx % 10 == 0:
            save_checkpoint(checkpoint)

    save_checkpoint(checkpoint)

    # ================= 宽表合并 =================
    if not all_dfs:
        print("\n未提取到任何有效数据，请检查输入目录或列名映射！")
        return

    print(f"\n成功提取 {len(all_dfs)} 张表，开始合并...")

    combined = pd.concat(all_dfs, axis=0, ignore_index=True, sort=False)

    # 列顺序：核心列在前，扩展列在后
    core_cols = [
        c for c in ['name', 'gender', 'id_card', 'level', 'sport',
                    'event', 'rank', 'team', 'province', 'year',
                    'ethnicity', 'phone', 'birthplace', 'apply_date',
                    'grant_date', 'cert_no', 'source_file']
        if c in combined.columns
    ]
    ext_cols = [c for c in combined.columns if c not in core_cols]
    combined = combined[core_cols + ext_cols]

    # 数据清洗
    if 'level' in combined.columns:
        combined['level'] = combined['level'].astype(str).str.strip()
        combined['level'] = combined['level'].replace({
            '一级': '一级运动员',
            '二级': '二级运动员',
            '三级': '三级运动员'
        })

    if 'gender' in combined.columns:
        combined['gender'] = combined['gender'].astype(str).str.strip()
        combined['gender'] = combined['gender'].apply(
            lambda x: x if x in ['男', '女'] else None
        )

    # 从 event 列补全省份和年份（如果之前没从文件名提取到）
    if 'event' in combined.columns:
        def extract_from_event(event, pattern_func):
            if pd.isna(event):
                return None
            return pattern_func(str(event))

        if 'province' in combined.columns:
            mask = combined['province'].isna()
            combined.loc[mask, 'province'] = combined.loc[mask, 'event'].apply(
                lambda x: extract_from_event(x, extract_province)
            )

        if 'year' in combined.columns:
            mask = combined['year'].isna()
            combined.loc[mask, 'year'] = combined.loc[mask, 'event'].apply(
                lambda x: extract_from_event(x, extract_year)
            )

    # 去重
    subset_cols = [
        c for c in ['name', 'id_card', 'level', 'event']
        if c in combined.columns
    ]
    if subset_cols:
        before = len(combined)
        combined = combined.drop_duplicates(subset=subset_cols, keep='first')
        print(f"去重：{before} → {len(combined)} 行")

    # 保存
    combined.to_excel(OUTPUT_FILE, index=False, engine='openpyxl')
    pd.DataFrame(logs).to_csv(LOG_FILE, index=False, encoding='utf-8-sig')

    print(f"\n{'='*50}")
    print(f"   处理完成")
    print(f"   总行数: {len(combined)}")
    print(f"   总列数: {len(combined.columns)}")
    print(f"   核心列 ({len(core_cols)}): {core_cols}")
    print(f"   扩展列 ({len(ext_cols)}): {ext_cols[:5]}... (共{len(ext_cols)}个)")
    print(f"   输出文件: {OUTPUT_FILE}")
    print(f"   处理日志: {LOG_FILE}")
    print(f"{'='*50}")


if __name__ == "__main__":
    main()